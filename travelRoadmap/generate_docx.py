import sys
import subprocess
import os

def install_and_import(import_name, package_name):
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package_name}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

install_and_import('docx', 'python-docx')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Time2Travel - UI Documentation Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run('Author: Sudeepa Santhanam\nRegister no.: 2547252')
author_run.bold = True
doc.add_page_break()

def add_placeholder(page_name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n[ PLACEHOLDER: Insert Screenshot of {page_name} Here ]\n')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True
    run.bold = True

doc.add_heading('1. Overview of the Application and its UI Goals', level=1)
doc.add_paragraph('Time2Travel is a comprehensive, interactive web application designed to offer smart travel planning within a specified budget. The primary UI goal of the application is to provide a seamless, premium, and highly engaging user journey from the moment of landing to the final itinerary review.')
doc.add_paragraph('The interface aims to evoke a sense of wanderlust and relaxation, achieved through a structured and visually rich layout that progressively guides the user through complex data inputs (e.g., stay selection, roadmap options, budgeting) without causing cognitive overload.')

doc.add_heading('2. Design Philosophy and Visual Direction', level=1)
doc.add_paragraph('The design philosophy revolves around Modern Elegance and Immersive Storytelling.')
doc.add_paragraph('• Color Palette: The UI employs a curated palette inspired by nature and travel. The central colors include deep "Brand" teal (#18465a), refreshing "Ocean" blues (#5699b3), sophisticated "Navy", and inviting "Warm" earthy tones (#65403a).')
doc.add_paragraph('• Glassmorphism: To maintain visual depth and a premium feel, "Glassmorphism" techniques (semi-transparent backgrounds with backdrop blur) are extensively used for cards, floating panels, and input elements.')
doc.add_paragraph('• Typography: The typography pairs the highly readable Inter font for body text and system UI with the elegant Playfair Display serif font for headings.')

doc.add_heading('3. Tools and Technologies Used', level=1)
doc.add_paragraph('• Frontend Framework: React 18 (with Vite) and TypeScript.')
doc.add_paragraph('• Routing: React Router DOM (v6) for seamless single-page application (SPA) transitions.')
doc.add_paragraph('• Styling Approach: Tailwind CSS for foundational utility-first styling, heavily customized to introduce complex animations, custom keyframes, and specific UI tokens.')
doc.add_paragraph('• Animations: Framer Motion for page-level transitions and component micro-interactions.')
doc.add_paragraph('• Maps Integration: React Leaflet (react-leaflet and leaflet) for the interactive Map View component.')
doc.add_paragraph('• Database & Real-time Integration: Supabase (PostgreSQL backend) is utilized to securely manage relational data and authentication.')

doc.add_heading('4. UI Architecture and Pages', level=1)
doc.add_paragraph('The UI architecture is modular and highly scalable. It consists of 14 primary pages orchestrated in an SPA model. Below are the visual representations for each page:')

pages = [
    'Landing Page',
    'How It Works Page',
    'Login Page',
    'Register Page',
    'Trip Planner Flow',
    'Stay Selection',
    'Roadmap Options',
    'Itinerary Page',
    'Map View',
    'Expense Breakdown',
    'Final Review',
    'User Profile',
    'Admin Dashboard',
    'Host Registration'
]

for p in pages:
    doc.add_heading(f'4.{pages.index(p)+1} {p}', level=2)
    add_placeholder(p)

doc.add_heading('5. Component Organization and Layout Strategy', level=1)
doc.add_paragraph('The application employs an atomic-like design strategy where UI components are broken down into logical units (e.g., Navbar, Footer, StarRating).')
doc.add_paragraph('Layouts predominantly rely on CSS Flexbox and CSS Grid, ensuring content remains structured horizontally for large screens while elegantly stacking vertically for smaller devices.')

doc.add_heading('6. Data Integration with Backend and How UI Reflects Live Data', level=1)
doc.add_paragraph('Data integration bridges the seamless UI with dynamic real-world inputs. The supabaseService.ts handles API calls to the remote Supabase database.')
doc.add_paragraph('• Live Syncing: React Context securely manages localized session states and form configurations, reflecting live updates gracefully.')
doc.add_paragraph('• State Feedback: The UI provides consistent state awareness using CSS shimmer loading states and real-time toast notifications.')

doc.add_heading('7. Use of Media Elements', level=1)
doc.add_paragraph('High-quality, themed images are consistently rendered within defined boundaries, specifically through image-cards with hover zoom capabilities. Dynamic floating "Decorative Blobs" provide subtle background visual mass. Gradients and interactive SVG icons (via lucide-react) act as visual anchors.')
add_placeholder('Component Details & Media Elements')

doc.add_heading('8. User Experience and Interaction Design', level=1)
doc.add_paragraph('User Interaction (UX) focuses heavily on immediate, positive feedback. Pages utilize sliding keyframes orchestrated by Framer Motion, eradicating the jarring visual effect of abrupt page loads. Buttons scale subtly when hovered over and emit a soft glowing shadow. Forms utilize Toaster notifications for instantaneous, non-blocking feedback during interactions.')

doc.add_heading('9. Responsiveness and Accessibility Considerations', level=1)
doc.add_paragraph('The application utilizes a mobile-first philosophy baked smoothly into Tailwind utility classes. Container widths and padding adjust fluidly to accommodate hand-held devices, tablets, and varied desktop monitor sizes. Clear visual hierarchy is maintained with distinct contrasts between text elements and the background, alongside semantic HTML structuring for accessibility.')
add_placeholder('Responsive Layout Comparison')

doc.add_heading('10. Design Improvements Made During Development', level=1)
doc.add_paragraph('During iterative development phases:')
doc.add_paragraph('• Transitioned from static solid-color blocks to complex gradient textures and backdrop blurring (Glassmorphism).')
doc.add_paragraph('• Optimized CSS transition durations to provide snappy yet smooth feedback.')
doc.add_paragraph('• Introduced Playfair Display heading typography to build a specialized, magazine-like reading experience.')

doc.add_heading('11. Justification of Design Choices', level=1)
doc.add_paragraph('Tailwind combined with Custom CSS delivers rapid prototyping while allowing granular control over complex, customized interactions. The React Router SPA inherently retains context and manages multi-step processes natively, protecting user-entered data safely. Glassmorphism conceptually matches the theme of travel and landscapes, enabling background imagery to permeate the layout cleanly.')

doc.add_heading('12. Workflow from Concept to Final UI Implementation', level=1)
doc.add_paragraph('1. Wireframing & Foundation: Configured root project settings, routing skeletons, and the global store.')
doc.add_paragraph('2. Design System Integration: Defined color palettes, fonts, custom utility classes, and custom animations.')
doc.add_paragraph('3. Component Drafting: Developed base components: Navbars, Footers, Input fields, and specialized cards.')
doc.add_paragraph('4. Page Assembly: Sequentially built interconnected pages leveraging reusable components.')
doc.add_paragraph('5. Backend Embellishment: Layered Supabase connectivity and context population.')
doc.add_paragraph('6. Refinement: Audited and optimized visualizations, transitions, and cross-screen responsiveness.')

output_path = os.path.join(os.getcwd(), 'UI_Documentation_Report.docx')
doc.save(output_path)
print(f'Successfully generated {output_path}')
